import json
import os
import random
import re
import sys
import tempfile
import threading
import time
import base64
import hashlib
import hmac
from collections import Counter
from collections import defaultdict
from datetime import datetime
from datetime import timedelta
from pathlib import Path
from typing import Any, Optional
from urllib.parse import quote_plus
from uuid import uuid4

import bcrypt
import requests
from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import desc, event, or_, text
from sqlalchemy.engine import Engine
from werkzeug.security import check_password_hash
from werkzeug.utils import secure_filename
try:
    from groq import Groq
except Exception:
    Groq = None
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


def _load_env_file(path: Path) -> None:
    """
    Lightweight .env loader so users don't need to export env vars in the terminal.
    - Supports KEY=VALUE lines (optional quotes).
    - Ignores blank lines and comments (# ...).
    - Does not overwrite already-set environment variables.
    """
    try:
        raw = path.read_text(encoding="utf-8")
    except Exception:
        return

    for line in raw.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        if not key or key in os.environ:
            continue
        value = value.strip().strip("'").strip('"')
        os.environ[key] = value


_load_env_file(PROJECT_ROOT / ".env")

from backend.course_service import get_courses_for_skills
from backend.resume_parser import ResumeParserError, extract_resume_text
from backend.resume_service import ResumeAnalysisService

MAX_MATCH_SCORE = 99
TOP_SEARCH_RESULTS = 500
TOP_MATCHED_JOBS = 50
TOP_RECOMMENDATIONS = TOP_MATCHED_JOBS
ADZUNA_API_URL_TEMPLATE = "https://api.adzuna.com/v1/api/jobs/in/search/{page}"
ADZUNA_RESULTS_PER_PAGE = 50
ADZUNA_TIMEOUT_SECONDS = 20
ADZUNA_MAX_PAGES_PER_REQUEST = 10
ADZUNA_REFRESH_PAGES = 10
CACHE_DURATION_SECONDS = 24 * 60 * 60
ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg"}
JWT_SECRET = os.getenv("VETERANHIRE_JWT_SECRET", "veteranhire-dev-secret")
JWT_EXPIRES_DAYS = 7

SKILL_KEYWORDS = {
    "administration",
    "advanced excel",
    "agile",
    "agile methodologies",
    "analytics",
    "application security",
    "as400",
    "asana",
    "asset management",
    "audit",
    "autocad",
    "aws",
    "azure",
    "bash",
    "bim",
    "budgeting",
    "business development",
    "coaching",
    "communication",
    "compliance",
    "coordination",
    "costing",
    "customer service",
    "data analysis",
    "data analytics",
    "data visualization",
    "devops",
    "discipline",
    "docker",
    "documentation",
    "electrical",
    "emergency response",
    "erp",
    "excel",
    "financial management",
    "financial modeling",
    "flask",
    "gcp",
    "gis",
    "go",
    "governance",
    "hr",
    "human resources",
    "incident response",
    "information security",
    "ios",
    "inventory management",
    "it support",
    "java",
    "javascript",
    "jira",
    "kubernetes",
    "leadership",
    "lean six sigma",
    "linux",
    "learning",
    "logistics",
    "machine learning",
    "maintenance",
    "marketing",
    "mentoring",
    "network security",
    "networking",
    "nist",
    "office 365",
    "operations",
    "penetration testing",
    "planning",
    "procurement",
    "process tracking",
    "project management",
    "python",
    "quality assurance",
    "quality control",
    "react",
    "reporting",
    "risk management",
    "safety",
    "sales",
    "sap",
    "sap s/4hana",
    "scrum",
    "scrum master",
    "scrum master certification",
    "security",
    "siem",
    "six sigma",
    "soc",
    "stakeholder communication",
    "stakeholder management",
    "supply chain management",
    "sql",
    "tableau",
    "technical writing",
    "threat hunting",
    "trello",
    "supply chain",
    "supervision",
    "team management",
    "technical support",
    "testing",
    "training",
    "transport",
    "troubleshooting",
    "udemy",
    "vendor management",
    "warehouse operations",
    "web development",
    "windows",
}

STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "to",
    "we",
    "with",
    "you",
    "your",
    "will",
    "our",
    "this",
    "their",
    "they",
    "role",
    "job",
    "required",
    "requirements",
    "experience",
    "work",
}

app = Flask(__name__)
CORS(app)

DEFAULT_DB_DIR = Path(os.getenv("VETERANHIRE_DB_DIR", Path(tempfile.gettempdir()) / "VeteranHire"))
DEFAULT_DB_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = Path(os.getenv("VETERANHIRE_DB_PATH", str(DEFAULT_DB_DIR / "veteranhire.db")))
UPLOAD_DIR = Path(os.getenv("VETERANHIRE_UPLOAD_DIR", str(DEFAULT_DB_DIR / "uploads")))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{DB_PATH}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "connect_args": {"check_same_thread": False, "timeout": 30},
}

db = SQLAlchemy(app)


@event.listens_for(Engine, "connect")
def set_sqlite_pragma(dbapi_connection, connection_record) -> None:
    try:
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA journal_mode=MEMORY")
        cursor.execute("PRAGMA synchronous=NORMAL")
        cursor.close()
    except Exception:
        pass


class User(db.Model):
    __tablename__ = "users"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(255), unique=True, nullable=False, index=True)
    password = db.Column(db.String(255), nullable=False)
    profile_pic = db.Column(db.String(400), nullable=True)
    phone = db.Column(db.String(25), nullable=True)
    location = db.Column(db.String(100), nullable=True)
    experience_years = db.Column(db.Integer, nullable=False, default=0)
    preferred_job_type = db.Column(db.String(30), nullable=True)
    education = db.Column(db.String(255), nullable=True)
    military_rank = db.Column(db.String(80), nullable=True)
    branch = db.Column(db.String(80), nullable=True)
    aadhaar_number = db.Column(db.String(20), nullable=True)
    role = db.Column(db.String(20), nullable=False, default="user")
    verification_status = db.Column(db.String(30), nullable=False, default="Pending")
    verification_document = db.Column(db.String(400), nullable=True)
    aadhaar_verification_document = db.Column(db.String(400), nullable=True)
    veteran_verification_document = db.Column(db.String(400), nullable=True)


class ResumeData(db.Model):
    __tablename__ = "resume_data"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=True, index=True)
    resume_text = db.Column(db.Text, nullable=False, index=True)
    extracted_skills = db.Column(db.Text, nullable=False)  # JSON string
    matched_jobs = db.Column(db.Text, nullable=False)  # JSON string
    skill_gap = db.Column(db.Text, nullable=False)  # JSON string
    match_score = db.Column(db.Integer, nullable=False, default=0)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)


class Job(db.Model):
    __tablename__ = "jobs"

    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False, index=True)
    company = db.Column(db.String(200), nullable=False, default="Unknown")
    description = db.Column(db.Text, nullable=False)
    skills = db.Column(db.Text, nullable=False)  # JSON string
    location = db.Column(db.String(100), nullable=False, default="India")
    category = db.Column(db.String(50), nullable=False, default="private", index=True)
    type = db.Column(db.String(50), nullable=False, default="private", index=True)
    salary_min = db.Column(db.Float, nullable=True)
    salary_max = db.Column(db.Float, nullable=True)
    experience_level = db.Column(db.String(20), nullable=False, default="0-2")
    link = db.Column(db.String(300), nullable=False, unique=True, index=True)


class Course(db.Model):
    __tablename__ = "courses"

    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(220), nullable=False, index=True)
    platform = db.Column(db.String(120), nullable=False, default="Coursera")
    level = db.Column(db.String(40), nullable=False, default="Beginner")
    duration = db.Column(db.String(60), nullable=False, default="4-6 weeks")
    rating = db.Column(db.Float, nullable=False, default=4.5)
    skill = db.Column(db.String(120), nullable=False, default="General")
    link = db.Column(db.String(400), nullable=False, unique=True, index=True)
    description = db.Column(db.Text, nullable=False, default="")
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)


class Application(db.Model):
    __tablename__ = "applications"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False, index=True)
    job_id = db.Column(db.Integer, nullable=False, index=True)
    status = db.Column(db.String(40), nullable=False, default="applied")
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow, index=True)

resume_analysis_service = ResumeAnalysisService(
    api_key=os.getenv("GEMINI_API_KEY"),
    model_name=os.getenv("GEMINI_MODEL", "gemini-2.0-flash"),
)
groq_api_key = os.getenv("GROQ_API_KEY", "").strip()
groq_client = Groq(api_key=groq_api_key) if Groq and groq_api_key else None
groq_model_primary = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile").strip()
groq_model_fallbacks = [
    "llama-3.1-8b-instant",
    "llama3-70b-8192",
    "mixtral-8x7b-32768",
]

_resume_cache_lock = threading.Lock()
_last_resume_result: dict[str, Any] = {}
_jobs_refresh_lock = threading.Lock()
_last_jobs_fetch_at = 0.0
_last_jobs_refresh_summary: dict[str, Any] = {
    "last_fetched_at": None,
    "added": 0,
    "fetched": 0,
    "keyword": "",
    "start_page": 1,
    "pages_requested": 0,
    "pages_failed": [],
}


def _normalize_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def _trim_text(value: Any, limit: int = 450) -> str:
    text_value = re.sub(r"\s+", " ", str(value or "").strip())
    if len(text_value) <= limit:
        return text_value
    return f"{text_value[:limit].rstrip()}..."


def _title_case_skill(skill: str) -> str:
    return " ".join(part.capitalize() for part in skill.split())


def _course_search_link(skill: str, provider: str) -> str:
    encoded_skill = quote_plus(str(skill or "").strip())
    provider_key = _normalize_text(provider)
    if provider_key == "coursera":
        return f"https://www.coursera.org/search?query={encoded_skill}"
    return f"https://www.udemy.com/courses/search/?q={encoded_skill}"


def _course_metadata() -> dict[str, str]:
    return {
        "rating": str(round(random.uniform(4.2, 4.8), 1)),
        "duration": random.choice(["4-6 weeks", "6-8 weeks", "2-3 months"]),
        "level": random.choice(["Beginner", "Intermediate"]),
    }


def _safe_json_loads(payload: str, fallback: Any) -> Any:
    try:
        return json.loads(payload)
    except Exception:
        return fallback


def _safe_positive_int(value: Any, default: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return int(default)
    return parsed if parsed > 0 else int(default)


def _b64url_encode(raw: bytes) -> str:
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def _b64url_decode(raw: str) -> bytes:
    value = str(raw or "").strip()
    pad = "=" * (-len(value) % 4)
    return base64.urlsafe_b64decode(value + pad)


def _jwt_hs256_encode(payload: dict[str, Any], secret: str) -> str:
    header = {"alg": "HS256", "typ": "JWT"}
    header_b64 = _b64url_encode(json.dumps(header, separators=(",", ":"), sort_keys=True).encode("utf-8"))
    payload_b64 = _b64url_encode(json.dumps(payload, separators=(",", ":"), sort_keys=True).encode("utf-8"))
    signing_input = f"{header_b64}.{payload_b64}".encode("ascii")
    signature = hmac.new(secret.encode("utf-8"), signing_input, hashlib.sha256).digest()
    sig_b64 = _b64url_encode(signature)
    return f"{header_b64}.{payload_b64}.{sig_b64}"


def _jwt_hs256_decode(token: str, secret: str) -> dict[str, Any]:
    parts = str(token or "").split(".")
    if len(parts) != 3:
        raise ValueError("invalid token")
    header_b64, payload_b64, sig_b64 = parts
    try:
        header = json.loads(_b64url_decode(header_b64))
        payload = json.loads(_b64url_decode(payload_b64))
    except Exception as exc:
        raise ValueError("invalid token") from exc

    if not isinstance(header, dict) or header.get("alg") != "HS256":
        raise ValueError("invalid token")

    signing_input = f"{header_b64}.{payload_b64}".encode("ascii")
    expected = hmac.new(secret.encode("utf-8"), signing_input, hashlib.sha256).digest()
    try:
        provided = _b64url_decode(sig_b64)
    except Exception as exc:
        raise ValueError("invalid token") from exc

    if not hmac.compare_digest(expected, provided):
        raise ValueError("invalid token")

    if not isinstance(payload, dict):
        raise ValueError("invalid token")

    exp = payload.get("exp")
    if exp is not None:
        try:
            exp_int = int(exp)
        except (TypeError, ValueError):
            raise ValueError("invalid token")
        if int(time.time()) >= exp_int:
            raise ValueError("expired token")

    return payload


def _hash_password(plain_password: str) -> str:
    hashed = bcrypt.hashpw(plain_password.encode("utf-8"), bcrypt.gensalt())
    return hashed.decode("utf-8")


def _verify_password(plain_password: str, stored_hash: str) -> bool:
    hash_value = str(stored_hash or "")
    if not hash_value:
        return False

    # Backward compatibility for existing werkzeug hashes.
    if hash_value.startswith("pbkdf2:") or hash_value.startswith("scrypt:"):
        return check_password_hash(hash_value, plain_password)

    try:
        return bcrypt.checkpw(plain_password.encode("utf-8"), hash_value.encode("utf-8"))
    except Exception:
        return False


def _issue_token(user: User) -> str:
    now = int(time.time())
    payload = {
        "user_id": int(user.id),
        "role": str(user.role or "user").strip().lower(),
        "iat": now,
        "exp": now + int(JWT_EXPIRES_DAYS * 24 * 60 * 60),
    }
    return _jwt_hs256_encode(payload, JWT_SECRET)


def _read_bearer_token() -> str:
    auth_header = str(request.headers.get("Authorization", "")).strip()
    if not auth_header.lower().startswith("bearer "):
        return ""
    return auth_header[7:].strip()


def _require_auth(required_role: Optional[str] = None) -> tuple[Optional[User], Optional[Any]]:
    token = _read_bearer_token()
    if not token:
        return None, (jsonify({"error": "Unauthorized"}), 401)

    try:
        payload = _jwt_hs256_decode(token, JWT_SECRET)
    except Exception:
        return None, (jsonify({"error": "Unauthorized"}), 401)

    user_id_raw = payload.get("user_id")
    if not isinstance(user_id_raw, int):
        try:
            user_id_raw = int(user_id_raw)
        except (TypeError, ValueError):
            return None, (jsonify({"error": "Unauthorized"}), 401)

    user = User.query.get(user_id_raw)
    if user is None:
        return None, (jsonify({"error": "Unauthorized"}), 401)

    if required_role and str(user.role or "").strip().lower() != str(required_role).strip().lower():
        return None, (jsonify({"error": "Unauthorized"}), 403)

    return user, None


def _tokenize_query(query: str) -> list[str]:
    return [token for token in re.split(r"[^a-z0-9]+", _normalize_text(query)) if token]


def _fit_label(score: int) -> str:
    if score >= 80:
        return "High Fit"
    if score >= 60:
        return "Medium Fit"
    return "Low Fit"


def _salary_to_lpa(value: Any) -> Optional[float]:
    try:
        amount = float(value)
    except (TypeError, ValueError):
        return None

    if amount <= 0:
        return None

    return round(amount / 100000.0, 2)


def _salary_range_label(min_lpa: Optional[float], max_lpa: Optional[float]) -> str:
    if min_lpa is None and max_lpa is None:
        return "Not specified"
    if min_lpa is None:
        return f"Up to {max_lpa:.1f} LPA"
    if max_lpa is None:
        return f"{min_lpa:.1f}+ LPA"
    return f"{min_lpa:.1f}-{max_lpa:.1f} LPA"


def _parse_salary_range_filter(raw_value: str) -> tuple[Optional[float], Optional[float]]:
    value = _normalize_text(raw_value)
    mapping = {
        "0-5l": (0.0, 5.0),
        "5-10l": (5.0, 10.0),
        "10-20l": (10.0, 20.0),
        "20+l": (20.0, None),
    }
    return mapping.get(value, (None, None))


def _extract_and_cache_resume_text(file_storage) -> str:
    file_bytes = file_storage.read()
    if not file_bytes:
        raise ValueError("The uploaded resume is empty.")

    try:
        return extract_resume_text(file_storage.filename, file_bytes)
    except ResumeParserError as exc:
        raise ValueError(str(exc)) from exc


def _normalize_skills(skills: Any) -> list[str]:
    if not isinstance(skills, list):
        return []

    normalized = []
    seen = set()
    for skill in skills:
        item = _normalize_text(skill)
        if item and item not in seen:
            seen.add(item)
            normalized.append(_title_case_skill(item))

    return normalized


def _extract_missing_skills(skill_gap: Any) -> list[str]:
    if not isinstance(skill_gap, dict):
        return []

    return _normalize_skills(skill_gap.get("missing_skills", []))


def _normalize_courses(courses: Any, missing_skills: list[str]) -> list[dict[str, str]]:
    normalized_courses = []
    seen = set()

    if isinstance(courses, list):
        for item in courses:
            if not isinstance(item, dict):
                continue

            name = str(item.get("name", "")).strip()
            provider = str(item.get("provider", item.get("platform", ""))).strip() or "Coursera"
            link = str(item.get("link", "")).strip()
            image = str(item.get("image", "")).strip()
            skill_query = str(item.get("skill", item.get("query", ""))).strip() or name

            if not name:
                continue
            provider_key = _normalize_text(provider)
            provider = "Coursera" if provider_key == "coursera" else "Udemy"
            link = _course_search_link(skill_query, provider)
            if any(domain in image.lower() for domain in ["placehold.co", "via.placeholder.com"]):
                image = ""

            key = (name.lower(), provider.lower())
            if key in seen:
                continue
            seen.add(key)

            course_meta = _course_metadata()
            normalized_courses.append(
                {
                    "name": name,
                    "provider": provider,
                    "skill": _title_case_skill(_normalize_text(skill_query)),
                    "link": link,
                    "image": image,
                    "rating": course_meta["rating"],
                    "duration": course_meta["duration"],
                    "level": course_meta["level"],
                }
            )

    if not normalized_courses and missing_skills:
        fallback = get_courses_for_skills(missing_skills)
        for course in fallback:
            name = str(course.get("name", "")).strip()
            provider = str(course.get("provider", "Coursera")).strip() or "Coursera"
            skill_query = str(course.get("skill", "")).strip() or name
            provider_key = _normalize_text(provider)
            provider = "Coursera" if provider_key == "coursera" else "Udemy"
            link = _course_search_link(skill_query, provider)
            if not name:
                continue
            course_meta = _course_metadata()
            normalized_courses.append(
                {
                    "name": name,
                    "provider": provider,
                    "skill": _title_case_skill(_normalize_text(skill_query)),
                    "link": link,
                    "image": "",
                    "rating": course_meta["rating"],
                    "duration": course_meta["duration"],
                    "level": course_meta["level"],
                }
            )

    return normalized_courses


def _normalize_roles(roles: Any) -> list[dict[str, Any]]:
    normalized_roles = []

    if not isinstance(roles, list):
        return normalized_roles

    for role in roles[:8]:
        if not isinstance(role, dict):
            continue

        title = str(role.get("title", "")).strip()
        if not title:
            continue

        try:
            score = int(round(float(role.get("match_score", 0))))
        except (TypeError, ValueError):
            score = 0

        normalized_roles.append(
            {
                "title": title,
                "match_score": max(0, min(MAX_MATCH_SCORE, score)),
                "skills_aligned": _normalize_skills(role.get("skills_aligned", [])),
            }
        )

    return normalized_roles


def _extract_skills_from_description(description: str) -> list[str]:
    text_value = _normalize_text(description)
    if not text_value:
        return []

    extracted = []
    for keyword in SKILL_KEYWORDS:
        if keyword in text_value:
            extracted.append(keyword)

    if extracted:
        # Prefer curated keywords for more stable matching across resumes/jobs.
        return [_title_case_skill(skill) for skill in sorted(set(extracted))][:18]

    tokens = re.findall(r"[a-z]{3,}", text_value)
    filtered = [token for token in tokens if token not in STOP_WORDS]
    common = [token for token, _ in Counter(filtered).most_common(12)]
    return [_title_case_skill(token) for token in common]


def _expand_skill_variants(skill_norm: str) -> set[str]:
    """Adds lightweight synonym variants to improve overlap without LLM calls."""
    variants: set[str] = set()
    if not skill_norm:
        return variants

    mapping = {
        "hr": {"human resources"},
        "human resources": {"hr"},
        "power bi": {"powerbi"},
        "powerbi": {"power bi"},
        "scrum master": {"scrum"},
        "scrum master certification": {"scrum", "scrum master"},
        "agile methodologies": {"agile"},
        "bim": {"building information modeling"},
        "building information modeling": {"bim"},
        "sap s/4hana": {"sap"},
    }
    if skill_norm in mapping:
        variants |= mapping[skill_norm]
    return variants


def _dedupe_skill_display(skills: list[str], limit: int = 14) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for raw in skills:
        norm = _normalize_text(raw)
        if not norm or norm in seen:
            continue
        seen.add(norm)
        output.append(_title_case_skill(norm))
        if len(output) >= limit:
            break
    return output


def _job_identity_key(title: Any, company: Any, location: Any) -> tuple[str, str, str]:
    return (
        _normalize_text(title),
        _normalize_text(company),
        _normalize_text(location),
    )


def _infer_job_category(title: str, description: str) -> str:
    blob = _normalize_text(f"{title} {description}")
    govt_markers = ["government", "ministry", "public sector", "psu", "railway", "defence", "army", "navy"]
    return "govt" if any(marker in blob for marker in govt_markers) else "private"


def _infer_experience_level(title: str, description: str) -> str:
    blob = _normalize_text(f"{title} {description}")
    year_hits = re.findall(r"(\d{1,2})\+?\s*(?:years|year|yrs|yr)", blob)
    if year_hits:
        max_years = max(int(hit) for hit in year_hits)
        if max_years >= 10:
            return "10+"
        if max_years >= 5:
            return "5-10"
        if max_years >= 2:
            return "2-5"
        return "0-2"

    if any(token in blob for token in ["senior", "lead", "manager", "director", "principal"]):
        return "10+"
    if any(token in blob for token in ["specialist", "experienced", "mid-level"]):
        return "5-10"
    if any(token in blob for token in ["junior", "entry", "associate"]):
        return "0-2"
    return "2-5"


def _serialize_job_row(job: Job) -> dict[str, Any]:
    skills = _safe_json_loads(job.skills, [])
    skills_list = skills if isinstance(skills, list) else []
    # Older DB rows may have too-few extracted skills. Enrich from title+description at read-time.
    if len(skills_list) < 6:
        derived = _extract_skills_from_description(f"{job.title} {job.description}")
        skills_list = _dedupe_skill_display([*skills_list, *derived], limit=18)
    category_value = str(getattr(job, "category", "") or getattr(job, "type", "private") or "private").strip().lower()
    return {
        "id": job.id,
        "title": job.title,
        "company": job.company,
        "description": job.description,
        "skills": skills_list,
        "location": job.location,
        "category": category_value,
        "type": category_value,
        "salary_min": job.salary_min,
        "salary_max": job.salary_max,
        "salary_range": _salary_range_label(job.salary_min, job.salary_max),
        "experience_level": job.experience_level,
        "link": job.link,
    }


def _search_score(job_payload: dict[str, Any], query_tokens: list[str]) -> float:
    if not query_tokens:
        return 0.0

    title = _normalize_text(job_payload.get("title", ""))
    description = _normalize_text(job_payload.get("description", ""))
    skills = [_normalize_text(skill) for skill in job_payload.get("skills", [])]

    keyword_hits = 0
    skill_hits = 0

    for token in query_tokens:
        if token in title:
            keyword_hits += 4
        if token in description:
            keyword_hits += 2
        if any(token in skill or skill in token for skill in skills):
            skill_hits += 3

    return float(keyword_hits + skill_hits)


def _match_jobs_rule_based(skills: list[str]) -> list[dict[str, Any]]:
    normalized_resume_skills: set[str] = set()
    for skill in skills:
        norm = _normalize_text(skill)
        if not norm:
            continue
        normalized_resume_skills.add(norm)
        normalized_resume_skills |= _expand_skill_variants(norm)
    if not normalized_resume_skills:
        return []

    ranked = []
    seen_job_identity: set[tuple[str, str, str]] = set()
    for row in Job.query.all():
        job_payload = _serialize_job_row(row)
        job_text = _normalize_text(f"{job_payload.get('title', '')} {job_payload.get('description', '')}")
        job_identity = _job_identity_key(
            job_payload.get("title", ""),
            job_payload.get("company", ""),
            job_payload.get("location", ""),
        )
        if job_identity in seen_job_identity:
            continue

        job_skill_tokens = {_normalize_text(skill) for skill in job_payload.get("skills", []) if _normalize_text(skill)}
        if not job_skill_tokens:
            continue

        overlap_set: set[str] = set(
            skill
            for skill in job_skill_tokens
            if any(skill in rs or rs in skill for rs in normalized_resume_skills)
        )
        # Also count resume skills that appear verbatim in the job text (helps when keyword extraction is sparse).
        if job_text:
            for rs in normalized_resume_skills:
                if len(rs) < 4:
                    continue
                if rs in job_text:
                    overlap_set.add(rs)

        overlap = sorted(overlap_set)
        if not overlap_set:
            continue
        missing = sorted(skill for skill in job_skill_tokens if skill not in overlap)

        # Score is overlap relative to the smaller of resume/job skill sets.
        denom = max(1, min(len(job_skill_tokens), len(normalized_resume_skills)))
        overlap_ratio = (len(overlap_set) / denom) * 100
        match_score = int(round(min(MAX_MATCH_SCORE, overlap_ratio)))

        seen_job_identity.add(job_identity)
        ranked.append(
            {
                "title": job_payload["title"],
                "company": job_payload["company"],
                "description": job_payload["description"],
                "type": job_payload["category"],
                "skills": job_payload["skills"],
                "link": job_payload["link"],
                "location": job_payload["location"],
                "salary_range": job_payload["salary_range"],
                "experience_level": job_payload["experience_level"],
                "match_score": match_score,
                "fit": _fit_label(match_score),
                "matched_skills": _dedupe_skill_display(overlap, limit=14),
                "missing_skills": _dedupe_skill_display(missing, limit=16),
            }
        )

    ranked.sort(key=lambda item: (item["match_score"], item["title"]), reverse=True)
    return ranked[:TOP_MATCHED_JOBS]


def _build_hybrid_response(ai_output: dict[str, Any]) -> dict[str, Any]:
    skills = _normalize_skills(ai_output.get("skills", []))
    roles = _normalize_roles(ai_output.get("roles", []))

    skill_gap_raw = ai_output.get("skill_gap", {})
    missing_skills = _extract_missing_skills(skill_gap_raw)
    courses = _normalize_courses(ai_output.get("courses", []), missing_skills)

    skill_gap = {
        "missing_skills": missing_skills,
        "recommended_courses": courses,
    }

    matched_jobs = _match_jobs_rule_based(skills)

    if not missing_skills and matched_jobs:
        missing_skills = matched_jobs[0].get("missing_skills", [])
        skill_gap["missing_skills"] = missing_skills
        if not courses:
            skill_gap["recommended_courses"] = _normalize_courses([], missing_skills)
            courses = skill_gap["recommended_courses"]

    top_job = matched_jobs[0] if matched_jobs else None
    top_match_score = int(top_job.get("match_score", 0)) if top_job else 0
    top_matched_skills = top_job.get("matched_skills", []) if top_job else []

    resume_data = {
        "skills": skills,
        "mapped_skills": skills,
        "experience": [],
    }

    return {
        "roles": roles,
        "jobs": matched_jobs,
        "skill_gap": skill_gap,
        "courses": courses,
        "match_score": min(MAX_MATCH_SCORE, top_match_score),
        "matched_skills": top_matched_skills,
        "missing_skills": skill_gap["missing_skills"],
        "recommended_jobs": matched_jobs,
        "resume_data": resume_data,
    }


def _store_resume_cache_record(
    user_id: Optional[int],
    resume_text: str,
    response_payload: dict[str, Any],
) -> None:
    extracted_payload = {
        "skills": response_payload.get("resume_data", {}).get("skills", []),
        "mapped_skills": response_payload.get("resume_data", {}).get("mapped_skills", []),
        "roles": response_payload.get("roles", []),
        "courses": response_payload.get("courses", []),
    }

    record = ResumeData(
        user_id=user_id,
        resume_text=resume_text,
        extracted_skills=json.dumps(extracted_payload),
        matched_jobs=json.dumps(response_payload.get("jobs", [])),
        skill_gap=json.dumps(response_payload.get("skill_gap", {"missing_skills": [], "recommended_courses": []})),
        match_score=int(response_payload.get("match_score", 0)),
    )
    db.session.add(record)
    db.session.commit()


def _load_cached_resume_record(record: ResumeData) -> dict[str, Any]:
    extracted = _safe_json_loads(
        record.extracted_skills,
        {"skills": [], "mapped_skills": [], "roles": [], "courses": []},
    )
    jobs = _safe_json_loads(record.matched_jobs, [])
    skill_gap = _safe_json_loads(record.skill_gap, {"missing_skills": [], "recommended_courses": []})

    roles = extracted.get("roles", []) if isinstance(extracted, dict) else []
    courses = extracted.get("courses", []) if isinstance(extracted, dict) else []
    if not courses:
        courses = skill_gap.get("recommended_courses", []) if isinstance(skill_gap, dict) else []

    resume_skills = extracted.get("skills", []) if isinstance(extracted, dict) else []
    mapped_skills = extracted.get("mapped_skills", resume_skills) if isinstance(extracted, dict) else []

    top_job = jobs[0] if jobs else {}
    # Heal old cached matches (no Gemini call; rule-based only).
    try:
        cached_matched = top_job.get("matched_skills", []) if isinstance(top_job, dict) else []
        if isinstance(resume_skills, list) and (
            not isinstance(jobs, list)
            or len(jobs) < min(20, TOP_MATCHED_JOBS)
            or (isinstance(cached_matched, list) and len(cached_matched) < 3)
        ):
            refreshed_jobs = _match_jobs_rule_based(_normalize_skills(resume_skills))
            if refreshed_jobs:
                jobs = refreshed_jobs
                top_job = jobs[0]
                record.matched_jobs = json.dumps(jobs)
                record.match_score = int(top_job.get("match_score", 0) or 0)
                db.session.commit()
    except Exception:
        pass

    return {
        "roles": roles,
        "jobs": jobs,
        "skill_gap": skill_gap,
        "courses": courses,
        "match_score": int(min(MAX_MATCH_SCORE, record.match_score or top_job.get("match_score", 0) or 0)),
        "matched_skills": top_job.get("matched_skills", []),
        "missing_skills": skill_gap.get("missing_skills", []),
        "recommended_jobs": jobs,
        "resume_data": {
            "skills": resume_skills,
            "mapped_skills": mapped_skills,
            "experience": [],
        },
        "cache_hit": True,
    }


def _build_user_courses(skill_gap: dict[str, Any]) -> list[dict[str, str]]:
    missing_skills = skill_gap.get("missing_skills", []) if isinstance(skill_gap, dict) else []
    response_rows = []

    for index, skill in enumerate(missing_skills):
        clean_skill = _title_case_skill(_normalize_text(skill))
        if not clean_skill:
            continue

        platform = "Coursera" if index % 2 == 0 else "Udemy"
        course_meta = _course_metadata()
        response_rows.append(
            {
                "skill": clean_skill,
                "course_name": f"{clean_skill} course",
                "platform": platform,
                "link": _course_search_link(clean_skill, platform),
                "rating": course_meta["rating"],
                "duration": course_meta["duration"],
                "level": course_meta["level"],
            }
        )

    return response_rows


def _build_chat_prompt(user_id: Optional[int], user_message: str) -> str:
    profile_summary = "No profile available."
    skill_summary = "Not available"
    experience_summary = "Not available"
    skill_gap_summary = "Not available"
    jobs_summary = "No jobs available."

    if user_id:
        user = User.query.get(user_id)
        if user is not None:
            profile_summary = (
                f"Name: {_trim_text(user.name, 80)} | "
                f"Location: {_trim_text(user.location or 'Not set', 60)} | "
                f"Experience (years): {int(user.experience_years or 0)} | "
                f"Preferred Job Type: {_trim_text(user.preferred_job_type or 'Not set', 20)}"
            )

        latest_resume = (
            ResumeData.query.filter_by(user_id=user_id)
            .order_by(desc(ResumeData.created_at))
            .first()
        )

        if latest_resume is not None:
            extracted = _safe_json_loads(latest_resume.extracted_skills, {})
            skill_gap = _safe_json_loads(latest_resume.skill_gap, {})
            matched_jobs = _safe_json_loads(latest_resume.matched_jobs, [])

            extracted_skills = extracted.get("skills", []) if isinstance(extracted, dict) else []
            if isinstance(extracted_skills, list) and extracted_skills:
                skill_summary = ", ".join(_normalize_skills(extracted_skills)[:15])

            experience_summary = (
                f"Latest resume processed on {latest_resume.created_at.strftime('%Y-%m-%d %H:%M UTC')}"
            )

            missing_skills = skill_gap.get("missing_skills", []) if isinstance(skill_gap, dict) else []
            if isinstance(missing_skills, list) and missing_skills:
                skill_gap_summary = ", ".join(_normalize_skills(missing_skills)[:12])

            if isinstance(matched_jobs, list) and matched_jobs:
                top_jobs = []
                for job in matched_jobs[:5]:
                    if not isinstance(job, dict):
                        continue
                    title = _trim_text(job.get("title", "Unknown Role"), 70)
                    company = _trim_text(job.get("company", "Unknown Company"), 50)
                    score = int(job.get("match_score", 0) or 0)
                    top_jobs.append(f"- {title} at {company} (Match: {score}%)")
                if top_jobs:
                    jobs_summary = "\n".join(top_jobs)

    return (
        "You are a career assistant for veterans.\n\n"
        "Response formatting rules:\n"
        "- Use plain text only; do not use markdown symbols such as ** or ##.\n"
        "- Write short paragraphs and bullet points where helpful.\n"
        "- Keep blank lines between sections for readability.\n"
        "- Keep output practical, supportive, and easy to scan.\n\n"
        "User Profile:\n"
        f"- {profile_summary}\n"
        f"- Skills: {skill_summary}\n"
        f"- Experience: {experience_summary}\n"
        f"- Skill Gap: {skill_gap_summary}\n\n"
        "Available Jobs:\n"
        f"{jobs_summary}\n\n"
        "User Question:\n"
        f"{_trim_text(user_message, 1200)}\n\n"
        "Give a helpful, structured answer."
    )


def _sanitize_assistant_reply(reply: str) -> str:
    text_value = str(reply or "").replace("\r\n", "\n").replace("\r", "\n")
    text_value = text_value.replace("**", "").replace("__", "").replace("`", "")
    text_value = re.sub(r"^\s{0,3}#{1,6}\s*", "", text_value, flags=re.MULTILINE)

    cleaned_lines: list[str] = []
    for raw_line in text_value.split("\n"):
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue

        # Normalize common list syntaxes to the required bullet style.
        if re.match(r"^(\d+[.)]|[-*])\s+", line):
            line = re.sub(r"^(\d+[.)]|[-*])\s+", "• ", line)
        elif line.startswith("•"):
            line = re.sub(r"^•\s*", "• ", line)

        line = line.replace("*", "")
        cleaned_lines.append(line)

    normalized = "\n".join(cleaned_lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
    return normalized


def _groq_model_candidates() -> list[str]:
    configured = [groq_model_primary, *groq_model_fallbacks]
    deduped: list[str] = []
    seen: set[str] = set()
    for model in configured:
        clean_model = str(model or "").strip()
        if clean_model and clean_model not in seen:
            deduped.append(clean_model)
            seen.add(clean_model)
    return deduped


SKILL_NORMALIZATION_MAP = {
    "pm": "Project Management",
    "ai": "Artificial Intelligence",
    "ml": "Machine Learning",
    "sql": "SQL",
    "hr": "Human Resources",
    "ops": "Operations Management",
}

ACTION_VERBS = {
    "achieved",
    "analyzed",
    "built",
    "coordinated",
    "created",
    "delivered",
    "developed",
    "drove",
    "executed",
    "implemented",
    "improved",
    "led",
    "managed",
    "optimized",
    "reduced",
    "streamlined",
}


def _title_case_text(value: str) -> str:
    words = [part for part in re.split(r"\s+", str(value or "").strip()) if part]
    return " ".join(word if word.isupper() else word.capitalize() for word in words)


def _split_csv_text(value: Any) -> list[str]:
    if value is None:
        return []
    tokens = []
    for piece in re.split(r"[,\n;]+", str(value)):
        clean = str(piece).strip()
        if clean:
            tokens.append(clean)
    return tokens


def _normalize_skill_label(skill: str) -> str:
    normalized_key = _normalize_text(skill)
    mapped = SKILL_NORMALIZATION_MAP.get(normalized_key, skill)
    return _title_case_text(mapped)


def _normalize_skill_list(skills_raw: Any) -> list[str]:
    values = skills_raw if isinstance(skills_raw, list) else _split_csv_text(skills_raw)
    normalized: list[str] = []
    seen: set[str] = set()
    for skill in values:
        label = _normalize_skill_label(str(skill))
        key = _normalize_text(label)
        if key and key not in seen:
            seen.add(key)
            normalized.append(label)
    return normalized


def _normalize_bullet_lines(value: Any) -> list[str]:
    parts = []
    for piece in re.split(r"[\n•]+", str(value or "")):
        clean = re.sub(r"^\s*[-*]\s*", "", piece).strip()
        if clean:
            parts.append(clean)

    if not parts:
        for sentence in re.split(r"[.;]+", str(value or "")):
            clean_sentence = sentence.strip()
            if clean_sentence:
                parts.append(clean_sentence)

    cleaned_bullets: list[str] = []
    seen: set[str] = set()
    for line in parts:
        line_text = line[0].upper() + line[1:] if line else line
        if not re.search(r"[.!?]$", line_text):
            line_text = f"{line_text}."
        key = _normalize_text(line_text)
        if key and key not in seen:
            seen.add(key)
            cleaned_bullets.append(line_text)
    return cleaned_bullets


def _ensure_skills_in_experience(experience_lines: list[str], skills: list[str]) -> list[str]:
    if not skills:
        return experience_lines

    joined_experience = _normalize_text(" ".join(experience_lines))
    missing_skills = [skill for skill in skills if _normalize_text(skill) not in joined_experience]
    augmented = list(experience_lines)
    for skill in missing_skills[:4]:
        augmented.append(f"Applied {skill} to improve operational outcomes and team performance.")
    return augmented


def _build_cv_payload(raw_payload: dict[str, Any]) -> dict[str, Any]:
    skills = _normalize_skill_list(raw_payload.get("skills", []))
    experience_lines = _normalize_bullet_lines(raw_payload.get("experience", ""))
    experience_lines = _ensure_skills_in_experience(experience_lines, skills)

    projects = _split_csv_text(raw_payload.get("projects", ""))
    awards = _split_csv_text(raw_payload.get("awards", ""))
    education = _split_csv_text(raw_payload.get("education", ""))

    return {
        "name": _trim_text(raw_payload.get("name", ""), 80),
        "role": _trim_text(raw_payload.get("role", ""), 80),
        "email": _trim_text(raw_payload.get("email", ""), 120),
        "phone": _trim_text(raw_payload.get("phone", ""), 40),
        "address": _trim_text(raw_payload.get("address", ""), 180),
        "education": education,
        "experience": experience_lines,
        "skills": skills,
        "projects": projects,
        "awards": awards,
        "summary": _trim_text(raw_payload.get("summary", ""), 900),
    }


def _calculate_ats_score(cv_payload: dict[str, Any]) -> dict[str, Any]:
    sections = {
        "name": bool(cv_payload.get("name")),
        "role": bool(cv_payload.get("role")),
        "contact": bool(cv_payload.get("email") and cv_payload.get("phone")),
        "summary": bool(cv_payload.get("summary")),
        "skills": bool(cv_payload.get("skills")),
        "experience": bool(cv_payload.get("experience")),
        "education": bool(cv_payload.get("education")),
        "projects": bool(cv_payload.get("projects")),
        "awards": bool(cv_payload.get("awards")),
    }

    completeness_ratio = sum(1 for present in sections.values() if present) / len(sections)
    completeness_score = int(round(completeness_ratio * 40))

    skill_count = len(cv_payload.get("skills", []))
    skill_score = min(35, skill_count * 4)

    experience_blob = _normalize_text(" ".join(cv_payload.get("experience", [])))
    verb_hits = sum(1 for verb in ACTION_VERBS if re.search(rf"\b{re.escape(verb)}\b", experience_blob))
    action_score = min(25, verb_hits * 5)

    total = min(100, completeness_score + skill_score + action_score)
    return {
        "score": total,
        "breakdown": {
            "section_completeness": completeness_score,
            "skill_relevance": skill_score,
            "action_verbs": action_score,
        },
    }


def _safe_resume_filename(name: str, extension: str) -> str:
    normalized = _normalize_text(name).replace(" ", "_") or "candidate"
    return f"{normalized}_ATS_Resume.{extension}"


def _generate_pdf_file(cv_payload: dict[str, Any]) -> Path:
    filename = _safe_resume_filename(cv_payload.get("name", ""), "pdf")
    output_path = UPLOAD_DIR / f"{uuid4().hex}_{filename}"

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        name="ATSHeading",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        name="ATSBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
    )
    name_style = ParagraphStyle(
        name="ATSName",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=20,
        spaceAfter=2,
    )
    role_style = ParagraphStyle(
        name="ATSRole",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        spaceAfter=8,
    )

    elements = [
        Paragraph(cv_payload.get("name", "Candidate"), name_style),
        Paragraph(cv_payload.get("role", ""), role_style),
        Paragraph(
            " | ".join(filter(None, [cv_payload.get("email", ""), cv_payload.get("phone", ""), cv_payload.get("address", "")])),
            body_style,
        ),
        Spacer(1, 6),
        Paragraph("Professional Summary", heading_style),
        Paragraph(cv_payload.get("summary", "N/A"), body_style),
        Paragraph("Skills", heading_style),
        Paragraph(", ".join(cv_payload.get("skills", [])) or "N/A", body_style),
        Paragraph("Experience", heading_style),
    ]

    for item in cv_payload.get("experience", []):
        elements.append(Paragraph(f"&bull; {item}", body_style))

    elements.extend([Paragraph("Education", heading_style)])
    for item in cv_payload.get("education", []):
        elements.append(Paragraph(f"&bull; {item}", body_style))

    elements.extend([Paragraph("Projects", heading_style)])
    for item in cv_payload.get("projects", []):
        elements.append(Paragraph(f"&bull; {item}", body_style))

    elements.extend([Paragraph("Certifications / Awards", heading_style)])
    for item in cv_payload.get("awards", []):
        elements.append(Paragraph(f"&bull; {item}", body_style))

    document.build(elements)
    return output_path


def _generate_docx_file(cv_payload: dict[str, Any]) -> Path:
    filename = _safe_resume_filename(cv_payload.get("name", ""), "docx")
    output_path = UPLOAD_DIR / f"{uuid4().hex}_{filename}"

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    doc.add_heading(cv_payload.get("name", "Candidate"), level=1)
    doc.add_paragraph(cv_payload.get("role", ""))
    doc.add_paragraph(" | ".join(filter(None, [cv_payload.get("email", ""), cv_payload.get("phone", ""), cv_payload.get("address", "")])))

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(cv_payload.get("summary", "N/A"))

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(cv_payload.get("skills", [])) or "N/A")

    doc.add_heading("Experience", level=2)
    for line in cv_payload.get("experience", []):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Education", level=2)
    for line in cv_payload.get("education", []):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Projects", level=2)
    for line in cv_payload.get("projects", []):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Certifications / Awards", level=2)
    for line in cv_payload.get("awards", []):
        doc.add_paragraph(line, style="List Bullet")

    doc.save(str(output_path))
    return output_path


def _build_course_response_from_ai_output(ai_output: dict[str, Any]) -> dict[str, Any]:
    skills = _normalize_skills(ai_output.get("skills", []))
    skill_gap_raw = ai_output.get("skill_gap", {})
    missing_skills = _extract_missing_skills(skill_gap_raw)
    courses = _normalize_courses(ai_output.get("courses", []), missing_skills)
    skill_gap = {
        "missing_skills": missing_skills,
        "recommended_courses": courses,
    }
    return {
        "skills": skills,
        "skill_gap": skill_gap,
        "courses": courses,
    }


def _build_course_response_from_record(record: ResumeData) -> dict[str, Any]:
    extracted = _safe_json_loads(record.extracted_skills, {"skills": [], "courses": []})
    skill_gap = _safe_json_loads(record.skill_gap, {"missing_skills": [], "recommended_courses": []})
    missing_skills = _extract_missing_skills(skill_gap)

    cached_courses: Any = []
    if isinstance(extracted, dict):
        cached_courses = extracted.get("courses", [])
    if not cached_courses and isinstance(skill_gap, dict):
        cached_courses = skill_gap.get("recommended_courses", [])

    courses = _normalize_courses(cached_courses, missing_skills)

    return {
        "skills": _normalize_skills(extracted.get("skills", []) if isinstance(extracted, dict) else []),
        "skill_gap": {
            "missing_skills": missing_skills,
            "recommended_courses": courses,
        },
        "courses": courses,
    }


def _ensure_users_schema() -> None:
    existing = db.session.execute(text("PRAGMA table_info(users)")).mappings().all()
    existing_cols = {row["name"] for row in existing}
    required_additions = {
        "profile_pic": "TEXT",
        "phone": "TEXT",
        "location": "TEXT",
        "experience_years": "INTEGER DEFAULT 0",
        "preferred_job_type": "TEXT",
        "education": "TEXT",
        "military_rank": "TEXT",
        "branch": "TEXT",
        "aadhaar_number": "TEXT",
        "role": "TEXT DEFAULT 'user'",
        "verification_status": "TEXT DEFAULT 'Pending'",
        "verification_document": "TEXT",
        "aadhaar_verification_document": "TEXT",
        "veteran_verification_document": "TEXT",
    }

    for col_name, ddl in required_additions.items():
        if col_name not in existing_cols:
            db.session.execute(text(f"ALTER TABLE users ADD COLUMN {col_name} {ddl}"))

    db.session.execute(
        text(
            "UPDATE users SET verification_status = CASE "
            "WHEN verification_status IS NULL OR TRIM(verification_status) = '' THEN 'Pending' "
            "WHEN lower(trim(verification_status)) = 'verified' THEN 'Verified' "
            "WHEN lower(trim(verification_status)) = 'rejected' THEN 'Rejected' "
            "ELSE 'Pending' END"
        )
    )
    db.session.execute(text("UPDATE users SET role = COALESCE(NULLIF(role, ''), 'user')"))
    admin_email = _normalize_text(os.getenv("VETERANHIRE_ADMIN_EMAIL", ""))
    if admin_email:
        db.session.execute(text("UPDATE users SET role='admin' WHERE lower(email)=:admin_email"), {"admin_email": admin_email})
    db.session.execute(text("UPDATE users SET experience_years = COALESCE(experience_years, 0)"))
    db.session.commit()


def _ensure_jobs_schema() -> None:
    existing = db.session.execute(text("PRAGMA table_info(jobs)")).mappings().all()
    existing_cols = {row["name"] for row in existing}
    required_additions = {
        "company": "TEXT DEFAULT 'Unknown'",
        "location": "TEXT DEFAULT 'India'",
        "category": "TEXT DEFAULT 'private'",
        "type": "TEXT DEFAULT 'private'",
        "salary_min": "REAL",
        "salary_max": "REAL",
        "experience_level": "TEXT DEFAULT '0-2'",
    }

    for col_name, ddl in required_additions.items():
        if col_name not in existing_cols:
            db.session.execute(text(f"ALTER TABLE jobs ADD COLUMN {col_name} {ddl}"))

    if "link" in existing_cols:
        try:
            db.session.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_jobs_link_unique ON jobs(link)"))
        except Exception:
            pass

    if "type" in existing_cols and "category" in existing_cols:
        db.session.execute(text("UPDATE jobs SET category = COALESCE(NULLIF(category, ''), type, 'private')"))
        db.session.execute(text("UPDATE jobs SET type = COALESCE(NULLIF(type, ''), category, 'private')"))
    db.session.execute(text("UPDATE jobs SET experience_level = COALESCE(NULLIF(experience_level, ''), '0-2')"))

    db.session.commit()


def _adzuna_credentials() -> tuple[str, str]:
    return os.getenv("ADZUNA_APP_ID", "").strip(), os.getenv("ADZUNA_APP_KEY", "").strip()


def _adzuna_url_for_page(page: int) -> str:
    safe_page = max(1, int(page))
    return ADZUNA_API_URL_TEMPLATE.format(page=safe_page)


def fetch_jobs_from_api(keyword: str = "", page: int = 1, pages: int = ADZUNA_REFRESH_PAGES) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    app_id, app_key = _adzuna_credentials()
    if not app_id or not app_key:
        raise RuntimeError("Adzuna credentials are missing. Set ADZUNA_APP_ID and ADZUNA_APP_KEY.")

    start_page = max(1, int(page))
    total_pages = max(1, min(int(pages), ADZUNA_MAX_PAGES_PER_REQUEST))
    all_jobs: list[dict[str, Any]] = []
    seen_links: set[str] = set()
    seen_identity: set[tuple[str, str, str]] = set()
    page_errors: list[dict[str, Any]] = []

    for current_page in range(start_page, start_page + total_pages):
        params = {
            "app_id": app_id,
            "app_key": app_key,
            "results_per_page": ADZUNA_RESULTS_PER_PAGE,
        }
        if keyword.strip():
            params["what"] = keyword.strip()

        try:
            response = requests.get(_adzuna_url_for_page(current_page), params=params, timeout=ADZUNA_TIMEOUT_SECONDS)
            response.raise_for_status()
            payload = response.json()
            results = payload.get("results", []) if isinstance(payload, dict) else []
            if not isinstance(results, list):
                results = []
        except Exception as exc:
            page_errors.append({"page": current_page, "error": str(exc)})
            continue

        if not results:
            break

        for item in results:
            if not isinstance(item, dict):
                continue

            title = str(item.get("title", "")).strip()
            description = str(item.get("description", "")).strip()
            link = str(item.get("redirect_url", "")).strip()
            company = str((item.get("company") or {}).get("display_name", "Unknown")).strip() or "Unknown"
            location = str((item.get("location") or {}).get("display_name", "India")).strip() or "India"

            safe_location = location[:100]
            identity_key = _job_identity_key(title, company, safe_location)
            if not title or not description or not link or link in seen_links or identity_key in seen_identity:
                continue

            skills = _extract_skills_from_description(description)
            category = _infer_job_category(title, description)
            salary_min = _salary_to_lpa(item.get("salary_min"))
            salary_max = _salary_to_lpa(item.get("salary_max"))
            experience_level = _infer_experience_level(title, description)
            safe_link = link[:300]
            seen_links.add(safe_link)
            seen_identity.add(identity_key)
            all_jobs.append(
                {
                    "title": title,
                    "company": company,
                    "description": description,
                    "skills": json.dumps(skills),
                    "location": safe_location,
                    "category": category,
                    "type": category,
                    "salary_min": salary_min,
                    "salary_max": salary_max,
                    "experience_level": experience_level,
                    "link": safe_link,
                }
            )

    return all_jobs, page_errors


def refresh_jobs(keyword: str = "", page: int = 1, pages: int = ADZUNA_REFRESH_PAGES) -> dict[str, Any]:
    global _last_jobs_fetch_at, _last_jobs_refresh_summary

    jobs_payload, page_errors = fetch_jobs_from_api(keyword=keyword, page=page, pages=pages)
    unique_jobs: list[dict[str, Any]] = []
    seen_identity: set[tuple[str, str, str]] = set()
    for job in jobs_payload:
        identity_key = _job_identity_key(
            job.get("title", ""),
            job.get("company", ""),
            job.get("location", ""),
        )
        if identity_key in seen_identity:
            continue
        seen_identity.add(identity_key)
        unique_jobs.append(job)

    Job.query.delete()
    db.session.commit()

    for job in unique_jobs:
        db.session.add(Job(**job))

    db.session.commit()

    result = {
        "added": len(unique_jobs),
        "fetched": len(jobs_payload),
        "keyword": keyword,
        "start_page": max(1, int(page)),
        "pages_requested": max(1, min(int(pages), ADZUNA_MAX_PAGES_PER_REQUEST)),
        "pages_failed": page_errors,
    }
    _last_jobs_fetch_at = time.time()
    _last_jobs_refresh_summary = {
        "last_fetched_at": datetime.utcnow().isoformat() + "Z",
        "added": result["added"],
        "fetched": result["fetched"],
        "keyword": result["keyword"],
        "start_page": result["start_page"],
        "pages_requested": result["pages_requested"],
        "pages_failed": result["pages_failed"],
    }
    return result


def get_jobs(force_refresh: bool = False) -> list[Job]:
    global _last_jobs_fetch_at

    should_refresh = force_refresh or Job.query.count() == 0 or (time.time() - _last_jobs_fetch_at > CACHE_DURATION_SECONDS)
    if should_refresh:
        with _jobs_refresh_lock:
            should_refresh_locked = (
                force_refresh
                or Job.query.count() == 0
                or (time.time() - _last_jobs_fetch_at > CACHE_DURATION_SECONDS)
            )
            if should_refresh_locked:
                try:
                    refresh_jobs(page=1, pages=ADZUNA_REFRESH_PAGES)
                    _last_jobs_fetch_at = time.time()
                except Exception:
                    # Keep serving existing DB data if refresh fails.
                    pass

    return Job.query.all()


def _auto_fetch_jobs_if_empty() -> None:
    get_jobs(force_refresh=False)


def _ensure_admin_user() -> None:
    admin_email = "admin@test.com"
    admin = User.query.filter_by(email=admin_email).first()
    if admin is None:
        admin = User(
            name="VeteranHire Admin",
            email=admin_email,
            password=_hash_password("admin123"),
            role="admin",
            verification_status="Verified",
        )
        db.session.add(admin)
        db.session.commit()
        return

    updated = False
    if str(admin.role or "").strip().lower() != "admin":
        admin.role = "admin"
        updated = True
    if not _verify_password("admin123", admin.password):
        admin.password = _hash_password("admin123")
        updated = True
    if updated:
        db.session.commit()


def _seed_demo_users(min_target: int = 40) -> None:
    seed_users = [
        ("user1@test.com", "User One", "user"),
        ("user2@test.com", "User Two", "user"),
        ("veteran@test.com", "Veteran User", "user"),
    ]
    for index in range(3, max(16, min_target + 1)):
        seed_users.append((f"demo{index}@test.com", f"Demo User {index}", "user"))

    for email, name, role in seed_users:
        normalized_email = _normalize_text(email)
        existing = User.query.filter_by(email=normalized_email).first()
        if existing is not None:
            # Ensure deterministic demo credentials.
            if not _verify_password("password123", existing.password):
                existing.password = _hash_password("password123")
            if str(existing.role or "").strip().lower() != role:
                existing.role = role
            continue

        db.session.add(
            User(
                name=name,
                email=normalized_email,
                password=_hash_password("password123"),
                role=role,
                verification_status=random.choice(["Pending", "Verified", "Rejected"]),
                experience_years=random.randint(0, 14),
                location=random.choice(["Delhi", "Mumbai", "Bengaluru", "Pune", "Hyderabad", "Chennai"]),
                preferred_job_type=random.choice(["govt", "private"]),
            )
        )
    db.session.commit()

    # If user table already has arbitrary rows, top up to target using generated demo users.
    current_total = User.query.count()
    next_index = max(17, current_total + 1)
    while current_total < min_target:
        generated_email = f"demo{next_index}@test.com"
        if User.query.filter_by(email=generated_email).first() is None:
            db.session.add(
                User(
                    name=f"Demo User {next_index}",
                    email=generated_email,
                    password=_hash_password("password123"),
                    role="user",
                    verification_status=random.choice(["Pending", "Verified", "Rejected"]),
                    experience_years=random.randint(0, 14),
                    location=random.choice(["Delhi", "Mumbai", "Bengaluru", "Pune", "Hyderabad", "Chennai"]),
                    preferred_job_type=random.choice(["govt", "private"]),
                )
            )
            db.session.commit()
            current_total = User.query.count()
        next_index += 1


def _seed_demo_jobs(min_target: int = 300) -> None:
    if Job.query.count() >= min_target:
        return

    titles = [
        "Logistics Coordinator",
        "Operations Supervisor",
        "Supply Chain Analyst",
        "Security Operations Lead",
        "Warehouse Manager",
        "Fleet Planner",
        "Training Coordinator",
        "Quality Assurance Specialist",
        "Procurement Executive",
        "Process Improvement Analyst",
    ]
    categories = ["govt", "private"]
    skills_pool = sorted(SKILL_KEYWORDS)

    existing_links = {row.link for row in Job.query.with_entities(Job.link).all()}
    attempts = 0
    max_attempts = max(500, min_target * 8)

    while Job.query.count() < min_target and attempts < max_attempts:
        attempts += 1
        title = random.choice(titles)
        company = random.choice(["Reliance", "Tata", "L&T", "DRDO Support", "RailServe", "SafeGrid", "OpsCore"])
        location = random.choice(["Delhi", "Mumbai", "Pune", "Bengaluru", "Hyderabad", "Chennai"])
        category = random.choice(categories)
        selected_skills = random.sample(skills_pool, k=min(5, len(skills_pool)))
        link = f"https://jobs.veteranhire.demo/{quote_plus(title.lower())}-{uuid4().hex[:10]}"
        if link in existing_links:
            continue
        existing_links.add(link)

        job = Job(
            title=title,
            company=company,
            description=f"{title} role focused on mission-critical delivery, disciplined execution, and team coordination.",
            skills=json.dumps([_title_case_skill(skill) for skill in selected_skills]),
            location=location,
            category=category,
            type=category,
            salary_min=random.choice([4.0, 6.0, 8.0, 10.0]),
            salary_max=random.choice([10.0, 12.0, 16.0, 20.0]),
            experience_level=random.choice(["0-2", "2-5", "5-10", "10+"]),
            link=link,
        )
        db.session.add(job)

        if attempts % 25 == 0:
            db.session.commit()

    db.session.commit()


def _seed_demo_courses(min_target: int = 60) -> None:
    if Course.query.count() >= min_target:
        return

    skill_seeds = [
        "Supply Chain",
        "Process Tracking",
        "Data Analysis",
        "Project Management",
        "Leadership",
        "Operations",
        "Security",
        "ERP",
        "Risk Management",
        "Vendor Management",
    ]
    platforms = ["Coursera", "Udemy"]
    levels = ["Beginner", "Intermediate"]

    existing_links = {row.link for row in Course.query.with_entities(Course.link).all()}
    attempts = 0
    max_attempts = max(500, min_target * 8)

    while Course.query.count() < min_target and attempts < max_attempts:
        attempts += 1
        skill = random.choice(skill_seeds)
        platform = random.choice(platforms)
        title = f"{skill} Mastery {Course.query.count() + 1}"
        link = f"https://www.google.com/search?q={quote_plus(skill + ' course ' + platform + ' ' + str(attempts))}"
        if link in existing_links:
            continue
        existing_links.add(link)

        db.session.add(
            Course(
                title=title,
                platform=platform,
                level=random.choice(levels),
                duration=random.choice(["4-6 weeks", "6-8 weeks", "8-10 weeks"]),
                rating=round(random.uniform(4.2, 4.9), 1),
                skill=skill,
                link=link,
                description=f"Hands-on {skill.lower()} coursework tailored for career transition.",
            )
        )

        if attempts % 25 == 0:
            db.session.commit()

    db.session.commit()


def _seed_demo_applications(min_target: int = 1200) -> None:
    user_ids = [row.id for row in User.query.with_entities(User.id).all()]
    job_ids = [row.id for row in Job.query.with_entities(Job.id).all()]
    if not user_ids or not job_ids:
        return
    max_pairs = len(user_ids) * len(job_ids)
    effective_target = min(min_target, max_pairs)
    if Application.query.count() >= effective_target:
        return

    existing_pairs = {(row.user_id, row.job_id) for row in Application.query.with_entities(Application.user_id, Application.job_id).all()}
    statuses = ["applied", "in_review", "shortlisted", "rejected"]
    attempts = 0
    max_attempts = max(2000, effective_target * 4)

    while Application.query.count() < effective_target and attempts < max_attempts:
        attempts += 1
        user_id = random.choice(user_ids)
        job_id = random.choice(job_ids)
        if (user_id, job_id) in existing_pairs:
            continue
        existing_pairs.add((user_id, job_id))

        days_ago = random.randint(0, 30)
        created_at = datetime.utcnow() - timedelta(days=days_ago, hours=random.randint(0, 23))
        db.session.add(
            Application(
                user_id=user_id,
                job_id=job_id,
                status=random.choice(statuses),
                created_at=created_at,
            )
        )
        if attempts % 50 == 0:
            db.session.commit()

    db.session.commit()


def _seed_demo_data(
    users_target: int = 40,
    jobs_target: int = 300,
    courses_target: int = 60,
    applications_target: int = 1200,
) -> None:
    _ensure_admin_user()
    _seed_demo_users(min_target=users_target)
    _seed_demo_jobs(min_target=jobs_target)
    _seed_demo_courses(min_target=courses_target)
    _seed_demo_applications(min_target=applications_target)


def _bootstrap_database() -> None:
    db.create_all()
    _ensure_users_schema()
    _ensure_jobs_schema()
    _auto_fetch_jobs_if_empty()
    _seed_demo_data()


def _relative_upload_path(path: Path) -> str:
    return f"/uploads/{path.name}"


def _public_upload_url(rel_path: Optional[str]) -> Optional[str]:
    return rel_path if rel_path else None


def _serialize_user(user: User) -> dict[str, Any]:
    return {
        "id": user.id,
        "name": user.name,
        "email": user.email,
        "phone": user.phone,
        "location": user.location,
        "experience_years": int(user.experience_years or 0),
        "preferred_job_type": user.preferred_job_type,
        "education": user.education,
        "military_rank": user.military_rank,
        "branch": user.branch,
        "aadhaar_number": user.aadhaar_number,
        "role": str(user.role or "user").strip().lower(),
        "verification_status": user.verification_status or "Pending",
        "profile_pic": _public_upload_url(user.profile_pic),
        "verification_document": _public_upload_url(user.verification_document),
        "aadhaar_verification_document": _public_upload_url(user.aadhaar_verification_document),
        "veteran_verification_document": _public_upload_url(user.veteran_verification_document),
    }


def _save_uploaded_file(file_storage, allowed_extensions: set[str]) -> str:
    filename = secure_filename(file_storage.filename or "")
    extension = Path(filename).suffix.lower()
    if not filename or extension not in allowed_extensions:
        raise ValueError("Unsupported file type.")

    unique_name = f"{uuid4().hex}{extension}"
    target_path = UPLOAD_DIR / unique_name
    file_storage.save(target_path)
    return _relative_upload_path(target_path)


@app.get("/uploads/<path:filename>")
def uploaded_file(filename: str):
    return send_from_directory(UPLOAD_DIR, filename)


@app.get("/health")
def health() -> Any:
    return jsonify(
        {
            "status": "ok",
            "gemini_configured": resume_analysis_service.rotator is not None,
            "groq_configured": groq_client is not None,
            "jobs_loaded": Job.query.count(),
            "db": "connected",
            "db_path": str(DB_PATH),
        }
    )


@app.post("/signup")
def signup() -> Any:
    payload = request.get_json(silent=True) or {}
    name = str(payload.get("name", "")).strip()
    email = _normalize_text(payload.get("email", ""))
    password = str(payload.get("password", ""))

    if not name or not email or not password:
        return jsonify({"error": "name, email, and password are required"}), 400

    if User.query.filter_by(email=email).first() is not None:
        return jsonify({"error": "Email already exists"}), 409

    admin_emails = {
        "admin@test.com",
        _normalize_text(os.getenv("VETERANHIRE_ADMIN_EMAIL", "")),
    }
    role = "admin" if email in admin_emails else "user"

    user = User(
        name=name,
        email=email,
        password=_hash_password(password),
        role=role,
        verification_status="Pending",
    )
    db.session.add(user)
    db.session.commit()

    token = _issue_token(user)
    return jsonify({"message": "Signup successful", "token": token, "user": _serialize_user(user)}), 201


@app.post("/login")
def login() -> Any:
    payload = request.get_json(silent=True) or {}
    email = _normalize_text(payload.get("email", ""))
    password = str(payload.get("password", ""))

    if not email or not password:
        return jsonify({"error": "email and password are required"}), 400

    user = User.query.filter_by(email=email).first()
    if user is None:
        return jsonify({"error": "User not found"}), 404
    if not _verify_password(password, user.password):
        return jsonify({"error": "Invalid password"}), 401

    token = _issue_token(user)
    return jsonify({"message": "Login successful", "token": token, "user": _serialize_user(user)})


@app.get("/profile/<int:user_id>")
def get_profile(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    return jsonify({"profile": _serialize_user(user)})


@app.put("/profile/<int:user_id>")
def update_profile(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    payload = request.get_json(silent=True) or {}

    email = payload.get("email")
    if email:
        normalized_email = _normalize_text(email)
        existing = User.query.filter(User.email == normalized_email, User.id != user_id).first()
        if existing:
            return jsonify({"error": "Email already in use"}), 409
        user.email = normalized_email

    if "name" in payload:
        user.name = str(payload.get("name", "")).strip() or user.name
    if "phone" in payload:
        user.phone = str(payload.get("phone", "")).strip() or None
    if "location" in payload:
        user.location = str(payload.get("location", "")).strip() or None
    if "experience_years" in payload:
        try:
            user.experience_years = max(0, int(payload.get("experience_years", 0)))
        except (TypeError, ValueError):
            user.experience_years = 0
    if "preferred_job_type" in payload:
        preferred = _normalize_text(payload.get("preferred_job_type", ""))
        user.preferred_job_type = preferred if preferred in {"govt", "private"} else None
    if "education" in payload:
        user.education = str(payload.get("education", "")).strip() or None
    if "military_rank" in payload:
        user.military_rank = str(payload.get("military_rank", "")).strip() or None
    if "branch" in payload:
        user.branch = str(payload.get("branch", "")).strip() or None
    if "aadhaar_number" in payload:
        aadhaar_raw = re.sub(r"\D+", "", str(payload.get("aadhaar_number", "")).strip())
        if aadhaar_raw and len(aadhaar_raw) != 12:
            return jsonify({"error": "aadhaar_number must contain exactly 12 digits"}), 400
        user.aadhaar_number = aadhaar_raw or None

    db.session.commit()
    return jsonify({"message": "Profile updated", "profile": _serialize_user(user)})


@app.post("/profile/<int:user_id>/profile_picture")
def upload_profile_picture(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    file_storage = request.files.get("file")
    if file_storage is None or not file_storage.filename:
        return jsonify({"error": "file is required"}), 400

    try:
        relative_path = _save_uploaded_file(file_storage, ALLOWED_IMAGE_EXTENSIONS)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400

    user.profile_pic = relative_path
    db.session.commit()
    return jsonify({"message": "Profile picture uploaded", "profile": _serialize_user(user)})


@app.post("/profile/<int:user_id>/verification")
def upload_verification(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    file_storage = request.files.get("file")
    if file_storage is None or not file_storage.filename:
        return jsonify({"error": "file is required"}), 400

    try:
        relative_path = _save_uploaded_file(file_storage, ALLOWED_DOCUMENT_EXTENSIONS)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400

    document_type = _normalize_text(request.form.get("document_type", "aadhaar"))
    aadhaar_number_raw = re.sub(r"\D+", "", str(request.form.get("aadhaar_number", "")).strip())
    if aadhaar_number_raw:
        if len(aadhaar_number_raw) != 12:
            return jsonify({"error": "aadhaar_number must contain exactly 12 digits"}), 400
        user.aadhaar_number = aadhaar_number_raw

    if document_type == "veteran":
        user.veteran_verification_document = relative_path
    else:
        user.aadhaar_verification_document = relative_path

    user.verification_document = relative_path
    user.verification_status = "Pending"
    db.session.commit()
    return jsonify(
        {
            "message": f"{'Veteran' if document_type == 'veteran' else 'Aadhaar'} verification document uploaded",
            "profile": _serialize_user(user),
        }
    )


@app.post("/profile/<int:user_id>/verification_status")
def set_verification_status(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    payload = request.get_json(silent=True) or {}
    status = str(payload.get("status", "")).strip().title()
    allowed = {"Pending", "Verified", "Rejected"}

    if status not in allowed:
        return jsonify({"error": "status must be one of: Pending, Verified, Rejected"}), 400

    user.verification_status = status
    db.session.commit()
    return jsonify({"message": "Verification status updated", "profile": _serialize_user(user)})


def _require_admin(admin_id_raw: Any) -> Optional[User]:
    auth_user, auth_error = _require_auth(required_role="admin")
    if auth_error is None and auth_user is not None:
        return auth_user

    # Backward-compatible fallback for legacy query parameter admin_id.
    if not admin_id_raw or not str(admin_id_raw).isdigit():
        return None
    admin_user = User.query.get(int(admin_id_raw))
    if admin_user is None:
        return None
    if str(admin_user.role or "user").strip().lower() != "admin":
        return None
    return admin_user


def _mask_aadhaar(value: Any) -> str:
    digits = re.sub(r"\D+", "", str(value or ""))
    if len(digits) != 12:
        return "Not submitted"
    return f"XXXX-XXXX-{digits[-4:]}"


def _collect_admin_stats_payload() -> dict[str, Any]:
    total_users = User.query.count()
    total_jobs = Job.query.count()
    total_applications = Application.query.count()

    active_cutoff = datetime.utcnow() - timedelta(days=30)
    active_user_ids = {
        row.user_id
        for row in Application.query.filter(Application.created_at >= active_cutoff)
        .with_entities(Application.user_id)
        .all()
    }
    active_users = len(active_user_ids)

    category_counter: dict[str, int] = defaultdict(int)
    for row in Job.query.with_entities(Job.category).all():
        category = str(row.category or "private").strip().lower() or "private"
        category_counter[category] += 1
    top_categories = [
        {"name": key, "count": value}
        for key, value in sorted(category_counter.items(), key=lambda item: item[1], reverse=True)[:6]
    ]

    recent_rows = (
        Application.query.order_by(desc(Application.created_at))
        .limit(12)
        .all()
    )
    recent_activity = []
    for row in recent_rows:
        user = User.query.get(row.user_id)
        job = Job.query.get(row.job_id)
        recent_activity.append(
            {
                "user": user.name if user else f"User {row.user_id}",
                "job_title": job.title if job else f"Job {row.job_id}",
                "status": row.status,
                "created_at": row.created_at.isoformat() if row.created_at else None,
            }
        )

    user_activity_counter: dict[str, int] = defaultdict(int)
    for days_ago in range(6, -1, -1):
        label = (datetime.utcnow() - timedelta(days=days_ago)).strftime("%Y-%m-%d")
        user_activity_counter[label] = 0
    for row in (
        Application.query.filter(Application.created_at >= datetime.utcnow() - timedelta(days=6))
        .with_entities(Application.created_at)
        .all()
    ):
        if row.created_at is None:
            continue
        key = row.created_at.strftime("%Y-%m-%d")
        if key in user_activity_counter:
            user_activity_counter[key] += 1
    user_activity = [{"date": key, "count": value} for key, value in user_activity_counter.items()]

    jobs_last_fetched_at = _last_jobs_refresh_summary.get("last_fetched_at")
    if not jobs_last_fetched_at and _last_jobs_fetch_at > 0:
        jobs_last_fetched_at = datetime.utcfromtimestamp(_last_jobs_fetch_at).isoformat() + "Z"

    return {
        "totalUsers": total_users,
        "totalJobs": total_jobs,
        "totalApplications": total_applications,
        "activeUsers": active_users,
        "topCategories": top_categories,
        "recentActivity": recent_activity,
        "userActivity": user_activity,
        "jobsLastFetchedAt": jobs_last_fetched_at,
        "jobsDedupedCount": total_jobs,
        "lastFetchAddedJobs": int(_last_jobs_refresh_summary.get("added") or 0),
        "lastFetchRawJobs": int(_last_jobs_refresh_summary.get("fetched") or 0),
    }


@app.get("/admin/users")
def admin_users() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Admin access required"}), 403

    rows = User.query.order_by(User.id.desc()).all()
    payload = []
    for row in rows:
        data = _serialize_user(row)
        payload.append(
            {
                "id": data["id"],
                "name": data["name"],
                "email": data["email"],
                "role": data.get("role", "user"),
                "aadhaar_masked": _mask_aadhaar(data.get("aadhaar_number")),
                "verification_status": data.get("verification_status", "Pending"),
            }
        )
    return jsonify({"users": payload})


@app.get("/admin/verification_queue")
def admin_verification_queue() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Admin access required"}), 403

    rows = (
        User.query.filter(User.verification_status.in_(["Pending", "Rejected"]))
        .order_by(User.id.desc())
        .all()
    )
    payload = []
    for row in rows:
        data = _serialize_user(row)
        payload.append(
            {
                "id": data["id"],
                "name": data["name"],
                "email": data["email"],
                "aadhaar_masked": _mask_aadhaar(data.get("aadhaar_number")),
                "verification_status": data.get("verification_status", "Pending"),
            }
        )
    return jsonify({"queue": payload})


@app.post("/admin/verification/<int:user_id>")
def admin_update_verification(user_id: int) -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Admin access required"}), 403

    payload = request.get_json(silent=True) or {}
    status = str(payload.get("status", "")).strip().title()
    if status not in {"Verified", "Rejected", "Pending"}:
        return jsonify({"error": "status must be one of: Pending, Verified, Rejected"}), 400

    target = User.query.get(user_id)
    if target is None:
        return jsonify({"error": "User not found"}), 404

    target.verification_status = status
    db.session.commit()
    return jsonify({"message": "Verification status updated", "user": _serialize_user(target)})


@app.get("/admin/stats")
def admin_stats() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403
    return jsonify(_collect_admin_stats_payload())


@app.post("/admin/seed_demo_data")
def admin_seed_demo_data() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403

    users_target = _safe_positive_int(request.args.get("users"), 40)
    jobs_target = _safe_positive_int(request.args.get("jobs"), 300)
    courses_target = _safe_positive_int(request.args.get("courses"), 60)
    applications_target = _safe_positive_int(request.args.get("applications"), 1200)

    before_counts = {
        "users": User.query.count(),
        "jobs": Job.query.count(),
        "courses": Course.query.count(),
        "applications": Application.query.count(),
    }

    _seed_demo_data(
        users_target=users_target,
        jobs_target=jobs_target,
        courses_target=courses_target,
        applications_target=applications_target,
    )

    after_counts = {
        "users": User.query.count(),
        "jobs": Job.query.count(),
        "courses": Course.query.count(),
        "applications": Application.query.count(),
    }

    return jsonify(
        {
            "message": "Demo data seeded successfully",
            "targets": {
                "users": users_target,
                "jobs": jobs_target,
                "courses": courses_target,
                "applications": applications_target,
            },
            "before": before_counts,
            "after": after_counts,
        }
    )


@app.get("/api/users")
def api_users() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403

    users = User.query.order_by(User.id.desc()).all()
    payload = []
    for user in users:
        serialized = _serialize_user(user)
        jobs_matched = Application.query.filter_by(user_id=user.id).count()
        payload.append(
            {
                "id": serialized["id"],
                "name": serialized["name"],
                "email": serialized["email"],
                "role": serialized["role"],
                "aadhaar_number": serialized.get("aadhaar_number"),
                "aadhaar_masked": _mask_aadhaar(serialized.get("aadhaar_number")),
                "verification_status": serialized.get("verification_status", "Pending"),
                "jobs_matched": jobs_matched,
            }
        )
    return jsonify({"users": payload})


@app.get("/api/verification")
def api_verification_queue() -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403

    rows = User.query.filter(User.verification_status.in_(["Pending", "Rejected"])).order_by(User.id.desc()).all()
    payload = []
    for user in rows:
        serialized = _serialize_user(user)
        payload.append(
            {
                "id": serialized["id"],
                "name": serialized["name"],
                "email": serialized["email"],
                "aadhaar_number": serialized.get("aadhaar_number"),
                "aadhaar_masked": _mask_aadhaar(serialized.get("aadhaar_number")),
                "verification_status": serialized.get("verification_status", "Pending"),
            }
        )

    return jsonify({"users": payload})


@app.post("/api/verify/<int:user_id>")
def api_verify_user(user_id: int) -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403

    payload = request.get_json(silent=True) or {}
    status = str(payload.get("status", "")).strip().title()
    if status not in {"Verified", "Rejected", "Pending"}:
        return jsonify({"error": "status must be one of: Pending, Verified, Rejected"}), 400

    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    user.verification_status = status
    db.session.commit()
    return jsonify({"message": "Verification updated", "user": _serialize_user(user)})


@app.get("/api/admin/user/<int:user_id>")
def api_admin_user_analytics(user_id: int) -> Any:
    admin_user = _require_admin(request.args.get("admin_id"))
    if admin_user is None:
        return jsonify({"error": "Unauthorized"}), 403

    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    applications = Application.query.filter_by(user_id=user_id).order_by(desc(Application.created_at)).all()
    latest_resume = ResumeData.query.filter_by(user_id=user_id).order_by(desc(ResumeData.created_at)).first()
    match_score = int(latest_resume.match_score) if latest_resume else 0

    status_counts: dict[str, int] = defaultdict(int)
    for item in applications:
        status_counts[str(item.status or "applied")] += 1

    activity = []
    for item in applications[:15]:
        job = Job.query.get(item.job_id)
        activity.append(
            {
                "job_title": job.title if job else f"Job {item.job_id}",
                "status": item.status,
                "created_at": item.created_at.isoformat() if item.created_at else None,
            }
        )

    return jsonify(
        {
            "user": {
                "id": user.id,
                "name": user.name,
                "email": user.email,
                "role": user.role,
                "verification_status": user.verification_status,
            },
            "stats": {
                "applications_count": len(applications),
                "jobs_matched": len({app.job_id for app in applications}),
                "match_score": max(0, min(MAX_MATCH_SCORE, match_score)),
            },
            "status_breakdown": [{"name": key, "value": value} for key, value in status_counts.items()],
            "recent_activity": activity,
        }
    )


@app.get("/fetch_jobs")
def fetch_jobs() -> Any:
    keyword = request.args.get("what", "")
    page_raw = request.args.get("page", "1")
    pages_raw = request.args.get("pages", "1")

    try:
        page = max(1, int(page_raw))
    except (TypeError, ValueError):
        page = 1

    try:
        pages = max(1, min(int(pages_raw), ADZUNA_MAX_PAGES_PER_REQUEST))
    except (TypeError, ValueError):
        pages = ADZUNA_REFRESH_PAGES

    try:
        result = refresh_jobs(keyword, page=page, pages=pages)
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 400
    except requests.RequestException as exc:
        return jsonify({"error": f"Adzuna API request failed: {exc}"}), 502
    except Exception as exc:
        return jsonify({"error": f"Failed to fetch jobs: {exc}"}), 500

    return jsonify(
        {
            "message": "Jobs fetched and stored successfully",
            "added": result["added"],
            "fetched": result["fetched"],
            "keyword": result["keyword"],
            "start_page": result["start_page"],
            "pages_requested": result["pages_requested"],
            "pages_failed": result["pages_failed"],
        }
    )


@app.get("/search_jobs")
def search_jobs() -> Any:
    get_jobs(force_refresh=False)

    query = request.args.get("q", "").strip()
    location = request.args.get("location", "").strip()
    salary_range = request.args.get("salary_range", "").strip()
    job_type = _normalize_text(request.args.get("job_type", "").strip())
    experience_level = request.args.get("experience_level", "").strip()

    jobs_query = Job.query
    if query:
        like_pattern = f"%{query}%"
        jobs_query = jobs_query.filter(
            or_(
                Job.title.ilike(like_pattern),
                Job.description.ilike(like_pattern),
                Job.skills.ilike(like_pattern),
            )
        )

    if location:
        jobs_query = jobs_query.filter(Job.location.ilike(f"%{location}%"))

    if job_type in {"govt", "private"}:
        jobs_query = jobs_query.filter(Job.category == job_type)

    if experience_level:
        jobs_query = jobs_query.filter(Job.experience_level == experience_level)

    range_min, range_max = _parse_salary_range_filter(salary_range)
    if range_min is not None:
        jobs_query = jobs_query.filter(or_(Job.salary_max.is_(None), Job.salary_max >= range_min))
    if range_max is not None:
        jobs_query = jobs_query.filter(or_(Job.salary_min.is_(None), Job.salary_min <= range_max))

    rows = jobs_query.order_by(Job.id.desc()).limit(TOP_SEARCH_RESULTS).all()
    payloads = [_serialize_job_row(row) for row in rows]

    if query:
        query_tokens = _tokenize_query(query)
        payloads.sort(key=lambda item: _search_score(item, query_tokens), reverse=True)

    results = []
    for payload in payloads:
        results.append(
            {
                "title": payload["title"],
                "company": payload["company"],
                "description": payload["description"],
                "type": payload["category"],
                "skills": payload["skills"],
                "location": payload["location"],
                "salary_range": payload["salary_range"],
                "experience_level": payload["experience_level"],
                "link": payload["link"],
            }
        )

    return jsonify(results)


@app.post("/recommend_jobs")
def recommend_jobs() -> Any:
    global _last_resume_result

    resume_file = request.files.get("resume_file")
    if resume_file is None or not resume_file.filename:
        return jsonify({"error": "resume_file is required"}), 400

    user_id_raw = request.form.get("user_id")
    user_id = int(user_id_raw) if user_id_raw and str(user_id_raw).isdigit() else None

    try:
        resume_text = _extract_and_cache_resume_text(resume_file)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"Resume parsing failed: {exc}"}), 502

    cached_record = ResumeData.query.filter_by(resume_text=resume_text).order_by(desc(ResumeData.created_at)).first()
    if cached_record is not None:
        response = _load_cached_resume_record(cached_record)
        with _resume_cache_lock:
            _last_resume_result = response
        return jsonify(response)

    try:
        ai_output = resume_analysis_service.analyze_resume_hybrid(resume_text)
        response = _build_hybrid_response(ai_output)
        _store_resume_cache_record(user_id=user_id, resume_text=resume_text, response_payload=response)
    except Exception as exc:
        return jsonify({"error": f"Resume processing failed: {exc}"}), 502

    with _resume_cache_lock:
        _last_resume_result = response

    return jsonify(response)


@app.post("/generate_courses")
def generate_courses() -> Any:
    resume_file = request.files.get("resume_file")

    payload = request.get_json(silent=True) or {}
    user_id_raw = request.form.get("user_id") if request.form else None
    if user_id_raw is None:
        user_id_raw = payload.get("user_id")

    user_id = int(user_id_raw) if user_id_raw and str(user_id_raw).isdigit() else None

    if resume_file is not None and resume_file.filename:
        try:
            resume_text = _extract_and_cache_resume_text(resume_file)
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except Exception as exc:
            return jsonify({"error": f"Resume parsing failed: {exc}"}), 502

        cached_record = ResumeData.query.filter_by(resume_text=resume_text).order_by(desc(ResumeData.created_at)).first()
        if cached_record is not None:
            response = _build_course_response_from_record(cached_record)
            response["cache_hit"] = True
            return jsonify(response)

        try:
            ai_output = resume_analysis_service.analyze_resume_hybrid(resume_text)
            response = _build_course_response_from_ai_output(ai_output)
        except Exception as exc:
            return jsonify({"error": f"Course generation failed: {exc}"}), 502

        _store_resume_cache_record(
            user_id=user_id,
            resume_text=resume_text,
            response_payload={
                "resume_data": {
                    "skills": response.get("skills", []),
                    "mapped_skills": response.get("skills", []),
                    "experience": [],
                },
                "roles": [],
                "jobs": [],
                "skill_gap": response.get("skill_gap", {"missing_skills": [], "recommended_courses": []}),
                "courses": response.get("courses", []),
                "match_score": 0,
            },
        )
        response["cache_hit"] = False
        return jsonify(response)

    if user_id is None:
        return jsonify({"error": "Provide either user_id or resume_file."}), 400

    latest_record = (
        ResumeData.query.filter_by(user_id=user_id)
        .order_by(desc(ResumeData.created_at))
        .first()
    )
    if latest_record is None:
        return jsonify({"error": "No stored resume found for this user. Please upload a resume."}), 404

    response = _build_course_response_from_record(latest_record)
    response["cache_hit"] = True
    return jsonify(response)


@app.get("/cv_prefill/<int:user_id>")
def cv_prefill(user_id: int) -> Any:
    user = User.query.get(user_id)
    if user is None:
        return jsonify({"error": "User not found"}), 404

    latest_resume = (
        ResumeData.query.filter_by(user_id=user_id)
        .order_by(desc(ResumeData.created_at))
        .first()
    )

    prefill: dict[str, Any] = {
        "name": user.name or "",
        "role": "",
        "email": user.email or "",
        "phone": user.phone or "",
        "address": user.location or "",
        "education": "",
        "experience": "",
        "skills": "",
        "projects": "",
        "awards": "",
        "summary": "",
    }

    if latest_resume is not None:
        extracted = _safe_json_loads(latest_resume.extracted_skills, {})
        matched_jobs = _safe_json_loads(latest_resume.matched_jobs, [])
        skills = extracted.get("skills", []) if isinstance(extracted, dict) else []
        if isinstance(skills, list):
            prefill["skills"] = ", ".join(_normalize_skill_list(skills))
        if isinstance(matched_jobs, list) and matched_jobs:
            top_role = matched_jobs[0]
            if isinstance(top_role, dict):
                prefill["role"] = str(top_role.get("title", "")).strip()

    if not prefill["role"]:
        prefill["role"] = "Operations Specialist"

    return jsonify({"prefill": prefill})


@app.post("/generate_pdf")
def generate_pdf() -> Any:
    payload = request.get_json(silent=True) or {}
    cv_payload = _build_cv_payload(payload)
    if not cv_payload.get("name"):
        return jsonify({"error": "name is required"}), 400

    ats_result = _calculate_ats_score(cv_payload)
    output_path = _generate_pdf_file(cv_payload)
    download_name = _safe_resume_filename(cv_payload.get("name", ""), "pdf")

    response = send_file(
        output_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=download_name,
    )
    response.headers["X-ATS-Score"] = str(ats_result["score"])
    return response


@app.post("/generate_docx")
def generate_docx() -> Any:
    payload = request.get_json(silent=True) or {}
    cv_payload = _build_cv_payload(payload)
    if not cv_payload.get("name"):
        return jsonify({"error": "name is required"}), 400

    ats_result = _calculate_ats_score(cv_payload)
    output_path = _generate_docx_file(cv_payload)
    download_name = _safe_resume_filename(cv_payload.get("name", ""), "docx")

    response = send_file(
        output_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=download_name,
    )
    response.headers["X-ATS-Score"] = str(ats_result["score"])
    return response


@app.post("/chat")
def chat() -> Any:
    payload = request.get_json(silent=True) or {}
    user_id_raw = payload.get("user_id")
    user_id = int(user_id_raw) if user_id_raw and str(user_id_raw).isdigit() else None
    message = str(payload.get("message", "")).strip()

    if not message:
        return jsonify({"error": "message is required"}), 400
    if groq_client is None:
        return jsonify({"error": "Groq is not configured. Set GROQ_API_KEY and install groq."}), 503

    full_prompt = _build_chat_prompt(user_id=user_id, user_message=message)

    response = None
    model_errors: list[str] = []
    for model_name in _groq_model_candidates():
        try:
            response = groq_client.chat.completions.create(
                model=model_name,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an AI career assistant for veterans. "
                            "Use clear plain text with short paragraphs and optional bullet lists. "
                            "Keep spacing between sections. "
                            "Do not use markdown symbols like ** or ##."
                        ),
                    },
                    {"role": "user", "content": full_prompt},
                ],
                temperature=0.4,
                max_tokens=600,
            )
            break
        except Exception as exc:
            model_errors.append(f"{model_name}: {exc}")

    if response is None:
        detail = " | ".join(model_errors[:3]) if model_errors else "No model response."
        return jsonify({"error": f"Chat request failed. {detail}"}), 502

    reply = ""
    choices = getattr(response, "choices", [])
    if choices:
        reply = str(getattr(choices[0].message, "content", "") or "").strip()
    if not reply:
        reply = "I'm sorry, I could not generate a response. Please try again."
    reply = _sanitize_assistant_reply(reply)

    return jsonify({"reply": reply})


@app.post("/career_guidance")
def career_guidance() -> Any:
    payload = request.get_json(silent=True) or {}
    action = _normalize_text(payload.get("action", ""))

    with _resume_cache_lock:
        current = _last_resume_result.copy()

    if not current:
        return jsonify({"error": "No processed resume found. Upload resume in Job Recommendations first."}), 400

    missing_skills = current.get("skill_gap", {}).get("missing_skills", [])

    if action == "improve_resume":
        tips = [
            "Add quantified achievements for each role (numbers, outcomes, timelines).",
            "Prioritize transferable civilian skills in the top summary section.",
            "Align experience bullets with the top missing skills from your job matches.",
        ]
        return jsonify({"title": "Improve My Resume", "items": tips})

    if action == "career_path":
        top_jobs = current.get("recommended_jobs", [])[:3]
        paths = [f"{job['title']} ({job['type'].upper()})" for job in top_jobs]
        if not paths:
            paths = ["Operations Coordinator", "Security Supervisor", "Training Specialist"]
        return jsonify({"title": "Suggested Career Path", "items": paths})

    if action == "skill_courses":
        return jsonify(
            {
                "title": "Skill-Based Course Recommendations",
                "missing_skills": missing_skills,
                "recommended_courses": current.get("skill_gap", {}).get("recommended_courses", []),
            }
        )

    return jsonify({"error": "Invalid action."}), 400


@app.get("/get_courses/<int:user_id>")
def get_courses(user_id: int) -> Any:
    latest_record = (
        ResumeData.query.filter_by(user_id=user_id)
        .order_by(desc(ResumeData.created_at))
        .first()
    )

    if latest_record is None:
        return jsonify({"error": "No resume data found for this user"}), 404

    skill_gap = _safe_json_loads(latest_record.skill_gap, {"missing_skills": [], "recommended_courses": []})
    course_rows = _build_user_courses(skill_gap)

    return jsonify(
        {
            "user_id": user_id,
            "missing_skills": skill_gap.get("missing_skills", []),
            "courses": course_rows,
        }
    )


@app.get("/")
def root() -> Any:
    return health()


with app.app_context():
    _bootstrap_database()


if __name__ == "__main__":
    with app.app_context():
        try:
            refresh_jobs(page=1, pages=ADZUNA_REFRESH_PAGES)
        except Exception:
            pass
    port = int(os.getenv("PORT", "8000"))
    app.run(host="0.0.0.0", port=port, debug=False)

